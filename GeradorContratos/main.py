from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

class GeradorDocumento():
    
    def __init__(self):
        self.contratante_nome = None
        self.contratante_rua = None
        self.contratante_cidade = None
        self.contratante_id = None

        self.contratado_nome = None
        self.contratado_rua = None
        self.contratado_cidade = None
        self.contratado_id = None


    def informacoes(self):
        print('*** Contratante ***')
        self.contratante_nome = input('Nome: ')
        self.contratante_rua = input('Rua: ')
        self.contratante_cidade = input('Cidade: ')
        self.contratante_id = input('Identificação: ')
        print()
        print('*** Contratado ***')
        self.contratado_nome = input('Nome: ')
        self.contratado_rua = input('Rua: ')
        self.contratado_cidade = input('Cidade: ')
        self.contratado_id = input('Identificação: ')


    def contrato_serviço(self):
        documento = Document()
        titulo_principal = documento.add_heading('Contrato de Prestação de Serviços', level=0).paragraph_format
        titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER


        titulo = documento.add_heading('Identificação das partes contratantes', level=1).paragraph_format
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        documento.add_paragraph('')


        paragrafo1 = documento.add_paragraph('')
        r1 = paragrafo1.add_run('Contratante: ').bold = True
        paragrafo1.add_run(f'{self.contratante_nome}, com sede em {self.contratante_rua} ({self.contratante_cidade}), inscrito sob o nº {self.contratante_id}.')
        documento.add_paragraph('')


        paragrafo2 = documento.add_paragraph('')
        r2 = paragrafo2.add_run('Contratado: ').bold = True
        paragrafo2.add_run(f'{self.contratado_nome}, com sede em {self.contratado_rua} ({self.contratado_cidade}), inscrito sob o nº {self.contratado_id}.')


        paragrafo3 = documento.add_paragraph('2. CLÁUSULA SEGUNDA - DO REGIME DE EXECUÇÃO\
        2.1. O serviço contratado será realizado por execução indireta, sob o regime de empreitada\
        por menor preço global.\
        3. CLÁUSULA TERCEIRA – DA FORMA DE PRESTAÇÃO DO SERVIÇO\
        3.1. Será considerada como unidade de pagamento a lauda completa com 1.000 (mil)\
        caracteres, eletronicamente contados pelo processador de textos no texto final, descontados\
        os espaços em branco, para a quantificação dos trabalhos.\
        3.2. O cálculo estimativo do número de laudas dar-se-á pelo uso do menu “ferramentas” e do.\
        3.3.3. “REGIME URGENTÍSSIMO” - produção acima de 20,01 (vinte vírgula zero um) laudas.')

        documento.add_paragraph('')
        documento.add_paragraph('')
        documento.add_paragraph('')
        documento.add_paragraph('')
        documento.add_paragraph('')

        assinatura_contratante = documento.add_paragraph('')
        assinatura_contratante.add_run('Assinatura contratante: ').bold = True
        assinatura_contratante.add_run('_________________________________________________')
        assinatura_contratante.alignment = WD_ALIGN_PARAGRAPH.CENTER

        documento.add_paragraph('')

        assinatura_contratado = documento.add_paragraph('')
        assinatura_contratado.add_run('Assinatura contratado: ').bold = True
        assinatura_contratado.add_run('_________________________________________________')
        assinatura_contratado.alignment = WD_ALIGN_PARAGRAPH.CENTER

        documento.save('contrato.docx')
        print('Seu contrato foi salvo na pasta!')
        return


GeradorDocumento = GeradorDocumento()
GeradorDocumento.informacoes()
GeradorDocumento.contrato_serviço()